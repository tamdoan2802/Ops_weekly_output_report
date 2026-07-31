from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import sys
import re
from docx import Document
import re
import json
import pandas as pd
import numpy as np
import unicodedata
from datetime import datetime, timedelta, time
from dateutil import parser as dateutil_parser

# Set output encoding to utf-8 to handle Vietnamese text properly
sys.stdout.reconfigure(encoding='utf-8')

# Global filters and configuration
EXCLUDE_USERS = ['huong huynh', 'unallocated', 'adrian', 'thuy ho', 'thuy  ho', 'long nguyen', 'tam doan']

def normalize_name(s):
    if not isinstance(s, str):
        return ""
    s = unicodedata.normalize('NFC', s.strip())
    s = " ".join(s.split())
    return s.lower()

def parse_date_column(series):
    def safe_parse(val):
        if pd.isna(val) or str(val).strip() in ('', 'nan', 'NaT'):
            return pd.NaT
        try:
            return dateutil_parser.parse(str(val))
        except Exception:
            return pd.NaT
    return series.apply(safe_parse)

def parse_time(s):
    try:
        return datetime.strptime(s.strip()[:5], "%H:%M").time()
    except:
        return None

def parse_detailed_cell(val):
    result = {"shifts": [], "leave_credit": 0.0, "leave_type": ""}
    if pd.isna(val) or str(val).strip() == "":
        return result
    
    val_str = str(val).strip()
    lines = val_str.split("\n")
    
    code_line_pattern = re.compile(r'^(.+?):\s*([\d.]+)\s*$')
    time_line_pattern = re.compile(r'^(\d{1,2}:\d{2}|-:-)\s*-\s*(\d{1,2}:\d{2}|-:-)$')
    
    current_shift = None
    for line in lines:
        line = line.strip()
        if not line or line.lower() == "cập nhật công":
            continue
        
        # Time line
        time_match = time_line_pattern.match(line)
        if time_match:
            ci_str = time_match.group(1)
            co_str = time_match.group(2)
            if current_shift is not None:
                current_shift["ci"] = parse_time(ci_str) if ci_str != "-:-" else None
                current_shift["co"] = parse_time(co_str) if co_str != "-:-" else None
            continue
            
        # Code line
        code_match = code_line_pattern.match(line)
        if code_match:
            code = code_match.group(1).strip()
            credit = float(code_match.group(2))
            
            # Check leave keywords
            leave_keywords = ["nghỉ phép", "nghỉ hưởng", "nghỉ ốm", "bhxh", "ốm", "thai sản", "bệnh", "nghỉ kết hôn", "tang", "hiếu", "phép"]
            is_leave = any(kw in code.lower() for kw in leave_keywords)
            
            if is_leave:
                result["leave_credit"] += credit
                if not result["leave_type"]:
                    result["leave_type"] = code
            else:
                current_shift = {"code": code, "work_credit": credit, "ci": None, "co": None}
                result["shifts"].append(current_shift)
                
    return result

def fmt_int(v):
    if pd.isna(v) or v == 0.0:
        return '-'
    try:
        return str(int(v))
    except ValueError:
        return '-'

def fmt_sqm(v):
    if pd.isna(v) or v == 0.0:
        return '-'
    try:
        return f"{v:.1f} m²"
    except ValueError:
        return '-'

def fmt_float(v, suffix='', decimals=1, show_plus=False):
    if pd.isna(v): return '-'
    try:
        val = float(v)
        if abs(val) < 1e-5: return '-'
        fmt_str = f"{{:.{decimals}f}}"
        res = fmt_str.format(val)
        if show_plus and val > 0:
            res = "+" + res
        return res + suffix
    except ValueError:
        return '-'

def get_task_flow_type(task_type_name):
    name_lower = str(task_type_name).lower()
    if 'design' in name_lower:
        return 'Design'
    elif 'check' in name_lower:
        return 'Check'
    elif 'review' in name_lower and 'fix' in name_lower:
        return 'R&F'
    elif 'train' in name_lower:
        return 'Training'
    return 'Other'

def calculate_week_ranges(target_end_date):
    """
    Calculate 5 consecutive weeks ending at target_end_date.
    Each week starts on Monday and ends on Sunday.
    """
    weeks = []
    # Find the Sunday of target_end_date's week if it isn't already a Sunday
    # weekday(): Monday is 0, Sunday is 6
    days_to_sunday = 6 - target_end_date.weekday()
    sunday = target_end_date + timedelta(days=days_to_sunday)
    
    for i in range(5):
        w_end = sunday - timedelta(weeks=i)
        w_start = w_end - timedelta(days=6)
        # Set times to start of day and end of day
        w_start = datetime(w_start.year, w_start.month, w_start.day, 0, 0, 0)
        w_end = datetime(w_end.year, w_end.month, w_end.day, 23, 59, 59)
        weeks.append((w_start, w_end))
        
    return list(reversed(weeks))

def get_kpi_target(company, kpi_config):
    config = kpi_config.get(company)
    if config is None:
        name_lower = company.lower()
        is_adhoc = any(k in name_lower for k in ["bm design", "all u want", "sto", "tilling timber", "wood shed", "rm property"])
        cust_type = "Ad-hoc" if is_adhoc else "Full-time"
        return cust_type, "Jobs", 0
    else:
        return config.get("Customer Type", "Full-time"), config.get("Expected Metric", "Jobs"), config.get("Daily Expected Output", 0)


def generate_docx_report(lines, docx_path):
    doc = Document()
    
    in_table = False
    table_lines = []
    
    def set_cell_bg(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def process_table():
        if not table_lines: return
        headers = [col.strip() for col in table_lines[0].strip('|').split('|')]
        data_rows = []
        for r_line in table_lines[1:]:
            clean_r = r_line.replace('-', '').replace(':', '').replace('|', '').strip()
            if not clean_r:
                continue
            data_rows.append([col.strip() for col in r_line.strip('|').split('|')])
            
        alignments = []
        if len(table_lines) > 1:
            sep_cols = [col.strip() for col in table_lines[1].strip('|').split('|')]
            for col in sep_cols:
                if col.startswith(':') and col.endswith(':'):
                    alignments.append('center')
                elif col.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Determine exact column widths to sum up to 19.0 cm for all table types
        col_widths = None
        num_cols = len(headers)
        if num_cols == 6:
            # Output trend or 6-col tables
            col_widths = [Cm(5.0), Cm(2.55), Cm(2.55), Cm(2.55), Cm(2.55), Cm(3.8)]
        elif num_cols == 8:
            # Member performance or Backlog tables
            first_header = headers[0].lower()
            if "team" in first_header or "client" in first_header or "khách" in first_header:
                col_widths = [Cm(4.3), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1), Cm(2.1)]
            else:
                col_widths = [Cm(3.8), Cm(2.1), Cm(1.9), Cm(2.1), Cm(1.9), Cm(2.4), Cm(2.4), Cm(2.4)]
        elif num_cols == 11:
            # Attendance table
            col_widths = [Cm(3.2), Cm(1.4), Cm(1.5), Cm(1.4), Cm(1.5), Cm(1.5), Cm(1.6), Cm(1.4), Cm(1.4), Cm(1.8), Cm(2.3)]
        elif num_cols == 5:
            # Task complexity table
            col_widths = [Cm(5.0), Cm(3.0), Cm(3.0), Cm(4.0), Cm(4.0)]
        else:
            # Generic fallback scaling to ~19cm
            w_each = 19.0 / num_cols
            col_widths = [Cm(w_each)] * num_cols

        # Format Header Row
        hdr_row = table.rows[0]
        hdr_row.height = Cm(1)
        hdr_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        hdr_cells = hdr_row.cells
        
        for i, header in enumerate(headers):
            if i < len(hdr_cells):
                cell = hdr_cells[i]
                if col_widths and i < len(col_widths):
                    cell.width = col_widths[i]
                set_cell_bg(cell, "061D47") # Dark Navy Blue
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.text = ""
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
                align = alignments[i] if i < len(alignments) else ('center' if i > 0 else 'left')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'right' else WD_ALIGN_PARAGRAPH.LEFT)
                
                run = p.add_run(header)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            
        # Format Data Rows
        for row_data in data_rows:
            row = table.add_row()
            row.height = Cm(1)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row_cells = row.cells
            
            first_col_text = row_data[0].strip() if len(row_data) > 0 else ""
            is_sub_item = first_col_text.startswith("•") or first_col_text.startswith("-") or first_col_text.startswith("↳")
            
            # Determine Row Shading
            if is_sub_item:
                bg = "FFFFFF"
            elif any(k in first_col_text for k in ["Completed Jobs", "Số Job hoàn thành", "Total Completed Jobs"]):
                bg = "98E398" # Light Green
            elif "Headcount" in first_col_text:
                bg = "EBA6EB" # Soft Lavender/Pink
            elif any(k in first_col_text for k in ["Units", "Area", "Layouts", "Sqm", "Designed", "Căn hộ", "Diện tích", "Bản vẽ"]):
                bg = "B8E4F7" # Light Sky Blue
            else:
                bg = "FFFFFF"
                
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    cell = row_cells[i]
                    if col_widths and i < len(col_widths):
                        cell.width = col_widths[i]
                    if bg != "FFFFFF":
                        set_cell_bg(cell, bg)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    
                    cell_clean = cell_data.replace('&nbsp;', ' ').strip()
                    p = cell.paragraphs[0]
                    p.text = ""
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    
                    align = alignments[i] if i < len(alignments) else ('center' if i > 0 else 'left')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else (WD_ALIGN_PARAGRAPH.RIGHT if align == 'right' else WD_ALIGN_PARAGRAPH.LEFT)
                    
                    parts = re.split(r'(\*\*.*?\*\*)', cell_clean)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r = p.add_run(part[2:-2])
                            r.bold = True
                        else:
                            r = p.add_run(part)
                            if is_sub_item:
                                r.italic = True
                            elif bg != "FFFFFF":
                                r.bold = True
                    
        table_lines.clear()

    for line in lines:
        line = line.strip()
        if not line:
            if in_table:
                process_table()
                in_table = False
            continue
            
        if line.startswith('|'):
            in_table = True
            table_lines.append(line)
        else:
            if in_table:
                process_table()
                in_table = False
                
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:].strip(), level=4)
            elif line.startswith('##### '):
                doc.add_heading(line[6:].strip(), level=5)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)

    if in_table:
        process_table()
        
    try:
        doc.save(docx_path)
        return docx_path
    except PermissionError:
        from datetime import datetime
        base, ext = os.path.splitext(docx_path)
        timestamp = datetime.now().strftime("%H%M%S")
        alt_path = f"{base}_{timestamp}{ext}"
        doc.save(alt_path)
        print(f"Warning: Target file '{docx_path}' was locked. Saved report to: {alt_path}")
        return alt_path


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Generate Weekly HR Output and Attendance Report')
    parser.add_argument('--start-date', type=str, help='Target week start date (YYYY-MM-DD)')
    parser.add_argument('--end-date', type=str, help='Target week end date (YYYY-MM-DD)')
    args = parser.parse_args()

    # Paths
    workload_path = r"G:\My Drive\Dữ liệu nhân sự\Workload\Construction Team\Report_Past Month.xlsx"
    attendance_path = r"G:\My Drive\Dữ liệu nhân sự\Data\Timesheet\HR_Fact_Attendance.xlsx"
    kpi_path = r"G:\My Drive\Dữ liệu nhân sự\.agents\skills\Ops_weekly_output_report\references\Customer KPI Mapping.json"
    output_dir = r"G:\My Drive\Dữ liệu nhân sự\.agents\skills\Ops_weekly_output_report\reports"
    
    if not os.path.exists(workload_path):
        print(f"Error: Workload file not found at '{workload_path}'")
        sys.exit(1)
        
    # Read files
    print("Loading Workload excel file...")
    try:
        jobs_df = pd.read_excel(workload_path, sheet_name="Jobs")
        tasks_df = pd.read_excel(workload_path, sheet_name="Tasks")
        emp_df = pd.read_excel(attendance_path, sheet_name="DIM_Employee")
    except Exception as e:
        print(f"Error reading source files: {e}")
        sys.exit(1)
        
    print("Parsing date columns...")
    jobs_df['Time Created'] = parse_date_column(jobs_df['Time Created'])
    jobs_df['Time Started'] = parse_date_column(jobs_df['Time Started'])
    jobs_df['Time Complete'] = parse_date_column(jobs_df['Time Complete'])
    
    tasks_df['timeCreated'] = parse_date_column(tasks_df['timeCreated'])
    tasks_df['timeCouldBeginAt'] = parse_date_column(tasks_df['timeCouldBeginAt'])
    tasks_df['timeOutcome'] = parse_date_column(tasks_df['timeOutcome'])

    # Determine reporting period (W0)
    if args.start_date and args.end_date:
        w0_start = datetime.strptime(args.start_date, '%Y-%m-%d')
        w0_end = datetime.strptime(args.end_date, '%Y-%m-%d')
        w0_end = datetime(w0_end.year, w0_end.month, w0_end.day, 23, 59, 59)
    else:
        today = datetime.now()
        current_monday = today - timedelta(days=today.weekday())
        w0_end = current_monday - timedelta(days=1)
        w0_start = w0_end - timedelta(days=6)
        w0_start = datetime(w0_start.year, w0_start.month, w0_start.day, 0, 0, 0)
        w0_end = datetime(w0_end.year, w0_end.month, w0_end.day, 23, 59, 59)
        
    print(f"Target week (W0): {w0_start.strftime('%d/%m/%Y')} – {w0_end.strftime('%d/%m/%Y')}")
    
    # Calculate the 5 weeks range (W-4 to W0)
    weeks = calculate_week_ranges(w0_end)
    print("Weeks to analyze:")
    for idx, (ws, we) in enumerate(weeks):
        print(f"  W-{4-idx}: {ws.strftime('%Y-%m-%d')} to {we.strftime('%Y-%m-%d')}")
        
    # Global Filters (Data_filtering_rule.md)
    jobs_df = jobs_df[jobs_df['Company'] != 'Test']
    tasks_df = tasks_df[tasks_df['jobId'].isin(jobs_df['Job #'])]
    tasks_df = tasks_df[~tasks_df['taskTypeName'].str.contains(r'Admin\s*-\s*QS|Admin\s*QS', case=False, na=False)]
    
    # Cutoff date (created after 01/04/2026)
    cutoff_date = pd.to_datetime("2026-04-01")
    jobs_df = jobs_df[jobs_df['Time Created'] > cutoff_date]
    tasks_df = tasks_df[tasks_df['timeCreated'] > cutoff_date]
    
    # Quick lookup mapping for jobs
    job_map = jobs_df.set_index('Job #').to_dict('index')

    # ── OVERRIDE JOB COMPLETION FOR RES ENGINEERING & PRIME DESIGN ───────────────
    print("Overriding Job Completion for RES Engineering & Prime Design based on Check tasks...")
    target_jobs = jobs_df[
        (jobs_df['Company'] == 'RES Engineering') |
        (jobs_df['Company'].str.contains('Prime Design', case=False, na=False))
    ]['Job #'].tolist()
    
    target_check_tasks = tasks_df[
        (tasks_df['jobId'].isin(target_jobs)) &
        (tasks_df['taskTypeName'].str.contains('Check', case=False, na=False)) &
        (tasks_df['Flow Task_status'].isin(['Complete', 'Completed'])) &
        (tasks_df['timeOutcome'].notna())
    ]
    
    if len(target_check_tasks) > 0:
        latest_checks = target_check_tasks.groupby('jobId')['timeOutcome'].max().reset_index()
        for _, row in latest_checks.iterrows():
            j_id = row['jobId']
            t_out = row['timeOutcome']
            
            job_mask = jobs_df['Job #'] == j_id
            jobs_df.loc[job_mask, 'Time Complete'] = t_out
            jobs_df.loc[job_mask, 'Status'] = 'Complete'
            
            if j_id in job_map:
                job_map[j_id]['Status'] = 'Complete'
                job_map[j_id]['Time Complete'] = t_out

    # Standardize IDs in DIM_Employee
    emp_df['EmployeeID'] = emp_df['EmployeeID'].astype(str).str.strip()
    
    # ── MAPPING WORKLOAD USERS TO TEAMS ──────────────────────────────────────────
    custom_username_to_en = {
        'adrian r': 'Adrian',
        'nam nguyen': 'Nguyễn Tuấn Nam',
        'phuoc vo': 'Quy Vo',
        'son ho': 'Hồ Thái Sơn',
        'thuy  ho': 'Thuy Ho',
    }
    
    workload_names = sorted(tasks_df['userName'].dropna().unique())
    user_to_team = {}
    workload_to_id = {}
    
    for name in workload_names:
        norm = normalize_name(name)
        if norm in EXCLUDE_USERS:
            user_to_team[name] = 'Excluded'
            continue
            
        mapped_target = custom_username_to_en.get(norm, name)
        emp_id = None
        if mapped_target == 'Nguyễn Tuấn Nam':
            emp_id = 'MTVN0090'
        elif mapped_target == 'Hồ Thái Sơn':
            emp_id = 'MTVN0085'
        elif mapped_target == 'Quy Vo':
            emp_id = 'MTVN0029'
        elif mapped_target == 'Thuy Ho':
            emp_id = 'MTVN0058'
        elif mapped_target == 'Adrian':
            emp_id = 'MTVN0059'
        else:
            match = emp_df[emp_df['FullNameEN'].apply(normalize_name) == normalize_name(mapped_target)]
            if len(match) > 0:
                emp_id = match.iloc[0]['EmployeeID']
            else:
                match_vn = emp_df[emp_df['FullNameVN'].apply(normalize_name) == normalize_name(mapped_target)]
                if len(match_vn) > 0:
                    emp_id = match_vn.iloc[0]['EmployeeID']
                    
        if emp_id:
            workload_to_id[name] = emp_id
            emp_row = emp_df[emp_df['EmployeeID'] == emp_id]
            if len(emp_row) > 0 and pd.notna(emp_row.iloc[0]['Team']) and str(emp_row.iloc[0]['Team']).strip() != '':
                user_to_team[name] = str(emp_row.iloc[0]['Team']).strip()
            else:
                user_to_team[name] = 'Unknown'
        else:
            user_to_team[name] = 'Unknown'

    # Helper function to attribute a completed Job to a Team
    def attribute_job_to_team(job_id):
        job_tasks = tasks_df[tasks_df['jobId'] == job_id]
        if len(job_tasks) == 0:
            return 'Unknown'
        
        # Design tasks first
        design_tasks = job_tasks[job_tasks['taskTypeName'].str.contains('Design', case=False, na=False)]
        for _, t in design_tasks.iterrows():
            user = t['userName']
            team = user_to_team.get(user)
            if team and team not in ('Unknown', 'Excluded'):
                return team
                
        # Check tasks
        check_tasks = job_tasks[job_tasks['taskTypeName'].str.contains('Check', case=False, na=False)]
        for _, t in check_tasks.iterrows():
            user = t['userName']
            team = user_to_team.get(user)
            if team and team not in ('Unknown', 'Excluded'):
                return team
                
        # Any task
        for _, t in job_tasks.iterrows():
            user = t['userName']
            team = user_to_team.get(user)
            if team and team not in ('Unknown', 'Excluded'):
                return team
                
        return 'Unknown'

    # ── METRIC 1: WEEKLY COMPLETED WORKLOAD TRENDS (5 WEEKS) BY TEAM ─────────────
    print("Calculating completed workload trend for the 5 weeks...")
    weekly_client_jobs = []
    weekly_client_jobs_detail = []
    
    for idx, (ws, we) in enumerate(weeks):
        week_label = f"W-{4-idx}"
        if idx == 4:
            week_label = "W0 (Current)"
            
        # Completed Jobs in this week
        comp_jobs_w = jobs_df[
            (jobs_df['Time Complete'] >= ws) &
            (jobs_df['Time Complete'] <= we) &
            (jobs_df['Status'] == 'Complete')
        ].copy()
        
        if len(comp_jobs_w) > 0:
            comp_jobs_w['Team'] = comp_jobs_w['Job #'].apply(attribute_job_to_team)
            comp_jobs_w = comp_jobs_w[~comp_jobs_w['Team'].isin(['Excluded', 'Unknown'])]
            
            comp_job_ids = comp_jobs_w['Job #'].tolist()
            design_tasks_w = tasks_df[
                (tasks_df['jobId'].isin(comp_job_ids)) &
                (tasks_df['Flow Task_status'].isin(['Complete', 'Completed'])) &
                (tasks_df['taskTypeName'].str.contains('Design', case=False, na=False))
            ].copy()
            
            # Numeric conversion
            for col in ['Designed Square Meter', 'Designed Layout']:
                design_tasks_w[col] = pd.to_numeric(design_tasks_w[col], errors='coerce').fillna(0)
            design_tasks_w['Dwelling Units'] = pd.to_numeric(design_tasks_w['Dwelling Units'], errors='coerce').fillna(0)
                
            design_tasks_w['Company'] = design_tasks_w['jobId'].apply(lambda x: job_map.get(x, {}).get('Company', 'Unknown'))
            design_tasks_w['Team'] = design_tasks_w['userName'].apply(lambda x: user_to_team.get(x, 'Unknown'))
            design_tasks_w = design_tasks_w[~design_tasks_w['Team'].isin(['Excluded', 'Unknown'])]
            
            # Group by Team and Company
            grouped_jobs = comp_jobs_w.groupby(['Team', 'Company']).size().reset_index(name='Jobs_Completed')
            grouped_vols = design_tasks_w.groupby(['Team', 'Company'])[['Dwelling Units', 'Designed Square Meter', 'Designed Layout']].sum().reset_index()
            grouped_headcounts = design_tasks_w.groupby(['Team', 'Company'])['userName'].nunique().reset_index(name='Headcount')
            
            merged_w = pd.merge(grouped_jobs, grouped_vols, on=['Team', 'Company'], how='outer').fillna(0)
            merged_w = pd.merge(merged_w, grouped_headcounts, on=['Team', 'Company'], how='outer').fillna(0)
            merged_w['Week'] = week_label
            merged_w['Week_Start'] = ws
            merged_w['Week_End'] = we
            weekly_client_jobs.append(merged_w)
            
            # Count jobs by Team, Company, Job Type
            grouped_jobs_detail = comp_jobs_w.groupby(['Team', 'Company', 'Job Type']).size().reset_index(name='Jobs_Completed')
            grouped_jobs_detail['Week'] = week_label
            weekly_client_jobs_detail.append(grouped_jobs_detail)
        else:
            empty_df = pd.DataFrame(columns=['Team', 'Company', 'Jobs_Completed', 'Dwelling Units', 'Designed Square Meter', 'Designed Layout', 'Headcount'])
            empty_df['Week'] = week_label
            empty_df['Week_Start'] = ws
            empty_df['Week_End'] = we
            weekly_client_jobs.append(empty_df)
            
            empty_detail = pd.DataFrame(columns=['Team', 'Company', 'Job Type', 'Jobs_Completed', 'Week'])
            weekly_client_jobs_detail.append(empty_detail)
            
    weekly_trend_df = pd.concat(weekly_client_jobs, ignore_index=True)
    weekly_trend_detail_df = pd.concat(weekly_client_jobs_detail, ignore_index=True)
    
    # ── METRIC 2: LOAD KPI TARGETS ────────────────────────────────────────────────
    print("Loading KPI mapping...")
    kpi_config = {}
    if os.path.exists(kpi_path):
        try:
            with open(kpi_path, "r", encoding="utf-8") as f:
                kpi_config = json.load(f)
        except Exception as e:
            print(f"Warning: Failed to load KPI config: {e}")
            
    # ── METRIC 3: AVERAGE LEADTIME FOR W0 JOBS BY TEAM ───────────────────────────
    print("Calculating Job Leadtimes for W0...")
    w0_completed_jobs = jobs_df[
        (jobs_df['Time Complete'] >= w0_start) &
        (jobs_df['Time Complete'] <= w0_end) &
        (jobs_df['Status'] == 'Complete')
    ].copy()
    
    if len(w0_completed_jobs) > 0:
        w0_completed_jobs['Leadtime_Days'] = (w0_completed_jobs['Time Complete'] - w0_completed_jobs['Time Started']).dt.total_seconds() / 86400.0
        w0_completed_jobs = w0_completed_jobs[w0_completed_jobs['Leadtime_Days'] >= 0]
        w0_completed_jobs['Team'] = w0_completed_jobs['Job #'].apply(attribute_job_to_team)
        w0_completed_jobs = w0_completed_jobs[~w0_completed_jobs['Team'].isin(['Excluded', 'Unknown'])]

    # ── METRIC 4: MEMBER WORKLOAD IN W0 ──────────────────────────────────────────
    print("Calculating Member Workload & Volume for W0...")
    w0_tasks = tasks_df[
        (tasks_df['timeOutcome'] >= w0_start) &
        (tasks_df['timeOutcome'] <= w0_end) &
        (tasks_df['Flow Task_status'].isin(['Complete', 'Completed']))
    ].copy()
    
    w0_tasks['norm_user'] = w0_tasks['userName'].apply(normalize_name)
    w0_tasks = w0_tasks[~w0_tasks['norm_user'].isin(EXCLUDE_USERS)]
    w0_tasks['FlowType'] = w0_tasks['taskTypeName'].apply(get_task_flow_type)
    
    for col in ['Dwelling Units', 'Designed Square Meter', 'Designed Layout']:
        w0_tasks[col] = pd.to_numeric(w0_tasks[col], errors='coerce').fillna(0)
        
    w0_tasks['Design_Units'] = w0_tasks.apply(lambda r: r['Dwelling Units'] if r['FlowType'] == 'Design' else 0.0, axis=1)
    w0_tasks['Design_Sqm'] = w0_tasks.apply(lambda r: r['Designed Square Meter'] if r['FlowType'] == 'Design' else 0.0, axis=1)
    w0_tasks['Design_Layouts'] = w0_tasks.apply(lambda r: r['Designed Layout'] if r['FlowType'] == 'Design' else 0.0, axis=1)
    
    member_workload = pd.DataFrame(columns=['userName', 'Design', 'Check', 'R&F', 'Training', 'Design_Units', 'Design_Sqm', 'Design_Layouts'])
    if len(w0_tasks) > 0:
        counts = w0_tasks.groupby(['userName', 'FlowType']).size().unstack(fill_value=0).reset_index()
        for col in ['Design', 'Check', 'R&F', 'Training']:
            if col not in counts.columns:
                counts[col] = 0
        vols = w0_tasks.groupby('userName')[['Design_Units', 'Design_Sqm', 'Design_Layouts']].sum().reset_index()
        member_workload = pd.merge(counts, vols, on='userName')
        member_workload = member_workload[['userName', 'Design', 'Check', 'R&F', 'Training', 'Design_Units', 'Design_Sqm', 'Design_Layouts']]
        
    # ── METRIC 5: INTRAWEEK COMPLETED TASKS IN W0 ────────────────────────────────
    print("Calculating Intraweek completed tasks...")
    w0_intraweek_tasks = tasks_df[
        (tasks_df['timeOutcome'] >= w0_start) &
        (tasks_df['timeOutcome'] <= w0_end) &
        (tasks_df['timeCouldBeginAt'] >= w0_start) &
        (tasks_df['timeCouldBeginAt'] <= w0_end) &
        (tasks_df['Flow Task_status'].isin(['Complete', 'Completed']))
    ].copy()
    
    w0_intraweek_tasks['norm_user'] = w0_intraweek_tasks['userName'].apply(normalize_name)
    w0_intraweek_tasks = w0_intraweek_tasks[~w0_intraweek_tasks['norm_user'].isin(EXCLUDE_USERS)]
    w0_intraweek_tasks['FlowType'] = w0_intraweek_tasks['taskTypeName'].apply(get_task_flow_type)
    
    intraweek_summary = pd.DataFrame(columns=['userName', 'Design', 'Check', 'R&F', 'Training'])
    if len(w0_intraweek_tasks) > 0:
        intraweek_summary = w0_intraweek_tasks.groupby(['userName', 'FlowType']).size().unstack(fill_value=0).reset_index()
        for col in ['Design', 'Check', 'R&F', 'Training']:
            if col not in intraweek_summary.columns:
                intraweek_summary[col] = 0
        intraweek_summary = intraweek_summary[['userName', 'Design', 'Check', 'R&F', 'Training']]

    # ── METRIC 5b: PER-CUSTOMER MEMBER WORKLOAD ───────────────────────────────────
    # Attach Company to w0_tasks via job_map
    w0_tasks['Company'] = w0_tasks['jobId'].map(lambda x: job_map.get(x, {}).get('Company', 'Unknown'))

    # Group counts per (userName, Company, FlowType)
    mc_counts_raw = (
        w0_tasks.groupby(['userName', 'Company', 'FlowType'])
        .size()
        .unstack(fill_value=0)
        .reset_index()
    )
    for col in ['Design', 'Check', 'R&F', 'Training']:
        if col not in mc_counts_raw.columns:
            mc_counts_raw[col] = 0

    # Group design volumes per (userName, Company)
    mc_vols = (
        w0_tasks.groupby(['userName', 'Company'])[['Design_Units', 'Design_Sqm', 'Design_Layouts']]
        .sum()
        .reset_index()
    )

    # Merge into one lookup df
    mc_workload_df = pd.merge(mc_counts_raw, mc_vols, on=['userName', 'Company'], how='outer').fillna(0)

    # ── METRIC 5c: PER-CUSTOMER MEMBER TASK COMPLEXITY ────────────────────────────
    # Build week_tasks early enough to use its Company column for the customer sections
    # (full computation happens again in Section 7 below, this is a lightweight subset)
    _wt = tasks_df.copy()
    _wt['timeOutcome_parsed'] = parse_date_column(_wt['timeOutcome'])
    _wt['timeOutcome_naive'] = _wt['timeOutcome_parsed'].apply(
        lambda x: x.replace(tzinfo=None) if pd.notna(x) else pd.NaT
    )
    _w0s_naive = w0_start.replace(tzinfo=None)
    _w0e_naive = w0_end.replace(tzinfo=None)

    _wt_week = _wt[
        (_wt['Flow Task_status'] == 'Complete') &
        (_wt['timeOutcome_naive'].notna()) &
        (_wt['timeOutcome_naive'] >= _w0s_naive) &
        (_wt['timeOutcome_naive'] <= _w0e_naive) &
        (_wt['investedSeconds'].notna()) &
        (_wt['investedSeconds'] > 0)
    ].copy()
    _wt_week['norm_user'] = _wt_week['userName'].apply(normalize_name)
    _wt_week = _wt_week[~_wt_week['norm_user'].isin(EXCLUDE_USERS)]
    _wt_week['Company'] = _wt_week['jobId'].map(lambda x: job_map.get(x, {}).get('Company', 'Unknown'))
    _wt_week['investedHours'] = _wt_week['investedSeconds'] / 3600.0

    def _get_prefix(s):
        s = str(s).strip()
        if s.startswith('Est.'): return 'Est'
        if s.startswith('Amd.'): return 'Amd'
        if s.startswith('Det.'): return 'Det'
        return 'Other'

    def _get_flow(s):
        sl = str(s).lower()
        if 'design' in sl: return 'Design'
        if 'check'  in sl: return 'Check'
        if 'review' in sl and 'fix' in sl: return 'R&F'
        return 'Other'

    _wt_week['ProjectPrefix'] = _wt_week['taskTypeName'].apply(_get_prefix)
    _wt_week['FlowType']      = _wt_week['taskTypeName'].apply(_get_flow)

    # Median investedHours per (userName, Company, ProjectPrefix, FlowType)
    mc_complexity = (
        _wt_week.groupby(['userName', 'Company', 'ProjectPrefix', 'FlowType'])['investedHours']
        .median()
    )  # MultiIndex Series; access with mc_complexity.get((user, company, prefix, flow), None)

    def fmt_hrs(val):
        """Format decimal hours → 'XhYYm' string."""
        if val is None or pd.isna(val) or val == 0:
            return '-'
        h = int(val)
        m = int(round((val - h) * 60))
        if h > 0 and m > 0: return f"{h}h{m:02d}m"
        if h > 0:            return f"{h}h"
        return f"{m}m"

    # ── METRIC 6: ATTENDANCE QUALITY INTEGRATION ─────────────────────────────────
    print("Loading Attendance & Leave data...")
    attendance_summary = pd.DataFrame()
    
    if os.path.exists(attendance_path):
        try:
            att_df = pd.read_excel(attendance_path, sheet_name="FACT_Attendance_Daily")
            leave_df = pd.read_excel(attendance_path, sheet_name="Req_Leave")
            ot_df = pd.read_excel(attendance_path, sheet_name="Req_OT")
            
            # Process OT for W0 (applies to both normal and fallback flows)
            ot_df['OT Date'] = pd.to_datetime(ot_df['OT Date'], errors='coerce')
            w0_ot = ot_df[
                (ot_df['Status'].astype(str).str.strip().str.lower() == 'đã duyệt') &
                (ot_df['OT Date'] >= w0_start) &
                (ot_df['OT Date'] <= w0_end)
            ].copy()
            w0_ot['OT_Hours'] = pd.to_numeric(w0_ot['OT_Hours'], errors='coerce').fillna(0.0)
            ot_dict = w0_ot.groupby('Employee_ID')['OT_Hours'].sum().to_dict()
            
            att_df['Date_Text'] = pd.to_datetime(att_df['Date_Text'], errors='coerce')
            w0_att = att_df[(att_df['Date_Text'] >= w0_start) & (att_df['Date_Text'] <= w0_end)].copy()
            
            if len(w0_att) > 0:
                w0_att['Type_of_Date'] = w0_att['Type of Date'].astype(str).str.strip()
                w0_att['Is_Full'] = w0_att['Type_of_Date'].apply(lambda x: 1 if x == 'FullWorkDay' else 0)
                w0_att['Is_Half'] = w0_att['Type_of_Date'].apply(lambda x: 1 if x == 'HalfWorkDay' else 0)
                
                w0_att['Working_Days'] = w0_att['Is_Full'] + w0_att['Is_Half']
                w0_att['Standard_Hours'] = w0_att['Is_Full'] * 8.0 + w0_att['Is_Half'] * 4.0
                w0_att['Actual_Hours'] = pd.to_numeric(w0_att['Số giờ làm việc thực tế'], errors='coerce').fillna(0.0)
                w0_att['Late_CI_Mins'] = pd.to_numeric(w0_att['Late_CheckIn (mins)'], errors='coerce').fillna(0.0)
                w0_att['Early_CO_Mins'] = pd.to_numeric(w0_att['Early_CO (mins)'], errors='coerce').fillna(0.0)
                
                w0_att['Is_Late_10'] = w0_att['Late_CI_Mins'].apply(lambda x: 1 if x > 10 else 0)
                w0_att['Is_Early_CO'] = w0_att['Early_CO_Mins'].apply(lambda x: 1 if x > 0 else 0)
                
                att_grouped = w0_att.groupby('Employee_ID').agg(
                    VN_Name=('Employee_Name', 'first'),
                    EN_Name=('DIM_Employee.FullNameEN', 'first'),
                    Working_Days=('Working_Days', 'sum'),
                    Standard_Hours=('Standard_Hours', 'sum'),
                    Actual_Hours=('Actual_Hours', 'sum'),
                    Late_Checkins=('Is_Late_10', 'sum'),
                    Early_Checkouts=('Is_Early_CO', 'sum')
                ).reset_index()
                
                leave_df['Leave_From'] = pd.to_datetime(leave_df['Leave_From'], errors='coerce')
                leave_df['Leave_To'] = pd.to_datetime(leave_df['Leave_To'], errors='coerce')
                
                w0_leaves = leave_df[
                    (leave_df['Status'].astype(str).str.strip().str.lower() == 'đã duyệt') &
                    (leave_df['Leave_From'] <= w0_end) &
                    (leave_df['Leave_To'] >= w0_start)
                ].copy()
                
                def calculate_w0_leave_days(row):
                    from_date = max(row['Leave_From'], w0_start).date()
                    to_date = min(row['Leave_To'], w0_end).date()
                    days = 0
                    curr = from_date
                    while curr <= to_date:
                        if curr.weekday() < 5:
                            days += 1
                        curr += timedelta(days=1)
                    orig_days = row['Leave_Days']
                    if orig_days < 1.0 and days > 0:
                        return orig_days
                    return min(days, orig_days)
                    
                w0_leaves['W0_Leave_Days'] = w0_leaves.apply(calculate_w0_leave_days, axis=1)
                w0_leaves['Leave_Type_Mapped'] = w0_leaves['Leave_Type'].apply(lambda x: 'Unpaid' if 'không lương' in str(x).lower() else 'Annual')
                
                leave_grouped = w0_leaves.groupby(['Employee_ID', 'Leave_Type_Mapped'])['W0_Leave_Days'].sum().unstack(fill_value=0.0).reset_index()
                for col in ['Annual', 'Unpaid']:
                    if col not in leave_grouped.columns:
                        leave_grouped[col] = 0.0
                        
                att_summary_combined = pd.merge(att_grouped, leave_grouped, on='Employee_ID', how='left').fillna(0.0)
                
                attendance_summary_rows = []
                for idx, row in att_summary_combined.iterrows():
                    emp_id = row['Employee_ID']
                    workload_user = '-'
                    for w_name, w_id in workload_to_id.items():
                        if w_id == emp_id:
                            workload_user = w_name
                            break
                    
                    attendance_summary_rows.append({
                        'EmployeeID': emp_id,
                        'VN_Name': row['VN_Name'],
                        'EN_Name': row['EN_Name'],
                        'Workload_User': workload_user,
                        'Working_Days': row.get('Working_Days', 0.0),
                        'Standard_Hours': row['Standard_Hours'],
                        'Actual_Hours': row['Actual_Hours'],
                        'OT_Hours': ot_dict.get(emp_id, 0.0),
                        'Late_Checkins': row['Late_Checkins'],
                        'Early_Checkouts': row['Early_Checkouts'],
                        'Annual_Leave': row['Annual'],
                        'Unpaid_Leave': row['Unpaid']
                    })
                attendance_summary = pd.DataFrame(attendance_summary_rows)
            else:
                print("FACT_Attendance_Daily has no records for target week. Falling back to detailed timesheet workbook...")
                if os.path.exists(ts_file):
                    shift_df = pd.read_excel(attendance_path, sheet_name="DIM_Shift")
                    shift_starts = {}
                    shift_ends = {}
                    for idx, row in shift_df.iterrows():
                        code = str(row['Shift_Code']).strip()
                        shift_starts[code] = row['Start_Time']
                        shift_ends[code] = row['End_Time']
                        
                    df_ts = pd.read_excel(ts_file, header=None, skiprows=10)
                    day_headers = [str(x).strip().split('.')[0] for x in df_ts.iloc[0].values]
                    
                    target_dates = []
                    curr_d = w0_start.date()
                    while curr_d <= w0_end.date():
                        target_dates.append(curr_d)
                        curr_d += timedelta(days=1)
                        
                    col_map = {}
                    for c in range(14, len(day_headers)):
                        day_str = day_headers[c]
                        if day_str and day_str != 'nan':
                            try:
                                d_val = int(float(day_str))
                                col_map[c] = d_val
                            except:
                                pass
                                
                    date_to_col = {}
                    for d in target_dates:
                        matched_c = None
                        for c, day_val in col_map.items():
                            if day_val == d.day:
                                if d.month == 6 and c < 24:
                                    matched_c = c
                                    break
                                elif d.month == 7 and c >= 24:
                                    matched_c = c
                                    break
                        if matched_c is not None:
                            date_to_col[d] = matched_c
                            
                    id_to_row = {}
                    for r_idx in range(1, len(df_ts)):
                        emp_id = str(df_ts.iloc[r_idx, 1]).strip()
                        if emp_id and emp_id != 'nan':
                            id_to_row[emp_id] = r_idx
                            
                    attendance_summary_rows = []
                    for name in workload_names:
                        norm = normalize_name(name)
                        if norm in EXCLUDE_USERS:
                            continue
                            
                        emp_id = workload_to_id.get(name)
                        if not emp_id or emp_id not in id_to_row:
                            continue
                            
                        r_idx = id_to_row[emp_id]
                        emp_row = df_ts.iloc[r_idx]
                        
                        total_working_days = 0.0
                        total_standard = 0.0
                        total_actual = 0.0
                        total_late = 0
                        total_early = 0
                        total_annual_leave = 0.0
                        total_unpaid_leave = 0.0
                        
                        for d, col in date_to_col.items():
                            cell_val = emp_row[col]
                            parsed = parse_detailed_cell(cell_val)
                            
                            n_full = sum(1 for s in parsed["shifts"] if s["work_credit"] == 1.0)
                            n_half = sum(1 for s in parsed["shifts"] if s["work_credit"] == 0.5)
                            
                            work_days = n_full + n_half
                            total_working_days += work_days
                            total_standard += n_full * 8.0 + n_half * 4.0
                            
                            if parsed["leave_credit"] > 0:
                                lt_mapped = 'Unpaid' if 'không lương' in str(parsed["leave_type"]).lower() else 'Annual'
                                if lt_mapped == 'Annual':
                                    total_annual_leave += parsed["leave_credit"]
                                else:
                                    total_unpaid_leave += parsed["leave_credit"]
                                    
                            for s in parsed["shifts"]:
                                if s["ci"] and s["co"]:
                                    t1 = datetime.combine(datetime.today(), s["ci"])
                                    t2 = datetime.combine(datetime.today(), s["co"])
                                    duration = (t2 - t1).total_seconds() / 3600.0
                                    if s["ci"] <= time(12, 0) and s["co"] >= time(13, 0):
                                        duration -= 1.0
                                    total_actual += max(0.0, duration)
                                    
                                    code = s["code"]
                                    if code in shift_starts:
                                        start_t = shift_starts[code]
                                        ci_delta = (t1 - datetime.combine(datetime.today(), start_t)).total_seconds() / 60.0
                                        if ci_delta > 10.0:
                                            total_late += 1
                                            
                                    if code in shift_ends:
                                        end_t = shift_ends[code]
                                        co_delta = (datetime.combine(datetime.today(), end_t) - t2).total_seconds() / 60.0
                                        if co_delta > 0.0:
                                            total_early += 1
                                            
                        attendance_summary_rows.append({
                            'EmployeeID': emp_id,
                            'VN_Name': emp_row[2],
                            'EN_Name': emp_row[3] if pd.notna(emp_row[3]) else '-',
                            'Workload_User': name,
                            'Working_Days': total_working_days,
                            'Standard_Hours': total_standard,
                            'Actual_Hours': total_actual,
                            'OT_Hours': ot_dict.get(emp_id, 0.0),
                            'Late_Checkins': total_late,
                            'Early_Checkouts': total_early,
                            'Annual_Leave': total_annual_leave,
                            'Unpaid_Leave': total_unpaid_leave
                        })
                    attendance_summary = pd.DataFrame(attendance_summary_rows)
                else:
                    print("Warning: Detailed timesheet workbook not found.")
        except Exception as e:
            print("Warning: Failed to process attendance quality data:", e)
            import traceback
            traceback.print_exc()

    # ── GENERATE MARKDOWN REPORT (TEAM → CUSTOMER → MEMBER) ─────────────────────
    print("Generating report text...")
    report_lines = []
    report_lines.append("# WEEKLY HR & WORKLOAD PERFORMANCE REPORT BY TEAM")
    report_lines.append(f"**Reporting Period:** {w0_start.strftime('%d/%m/%Y')} – {w0_end.strftime('%d/%m/%Y')} (Week {w0_start.isocalendar()[1]} / {w0_start.year})")
    report_lines.append(f"**Report Date:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    report_lines.append("\n---\n")


    active_teams = ['Frame & Truss', 'Drafting', 'Estimating', 'Engineered Wood Products (EWP)', 'Cleveland']
    week_labels  = ["W-4", "W-3", "W-2", "W-1", "W0 (Current)"]

    for t_idx, team_name in enumerate(active_teams, 1):
        report_lines.append(f"# {t_idx}. TEAM: {team_name.upper()}")
        report_lines.append(f"Workload volume analysis for Clients (Customers) and member productivity for Team **{team_name}**.")
        report_lines.append("")

        team_members = [k for k, v in user_to_team.items() if v == team_name]
        team_trend   = weekly_trend_df[weekly_trend_df['Team'] == team_name]
        team_detail  = weekly_trend_detail_df[weekly_trend_detail_df['Team'] == team_name]

        # Build ordered customer list: sort by W0 jobs desc (full-time first, then ad-hoc)
        all_companies_5w = sorted(team_trend['Company'].dropna().unique().tolist())
        w0_jobs_per_comp = (
            team_trend[team_trend['Week'] == 'W0 (Current)']
            .groupby('Company')['Jobs_Completed'].sum()
            .to_dict()
        )

        def cust_sort_key(c):
            ctype, _, _ = get_kpi_target(c, kpi_config)
            tier = 0 if ctype == 'Full-time' else 1
            return (tier, -w0_jobs_per_comp.get(c, 0), c)

        ordered_companies = sorted(all_companies_5w, key=cust_sort_key)

        # ── PER-CUSTOMER SECTIONS ─────────────────────────────────────────────
        for c_idx, company in enumerate(ordered_companies, 1):
            cust_type, metric, daily_target = get_kpi_target(company, kpi_config)
            weekly_target = daily_target * 5

            report_lines.append(f"## {t_idx}.{c_idx}. {company}")
            report_lines.append(
                f"**Client Type:** {cust_type}  |  "
                f"**Weekly Target:** {weekly_target if weekly_target > 0 else 'N/A'} {metric}"
            )
            report_lines.append("")

            # 5-week trend table for this company
            comp_trend = team_trend[team_trend['Company'] == company]
            comp_detail = team_detail[team_detail['Company'] == company]

            jobs_val    = []
            units_val   = []
            sqm_val     = []
            layouts_val = []
            headcounts_val = []
            for wl in week_labels:
                sub = comp_trend[comp_trend['Week'] == wl]
                jobs_val.append(sub['Jobs_Completed'].sum())
                units_val.append(sub['Dwelling Units'].sum())
                sqm_val.append(sub['Designed Square Meter'].sum())
                layouts_val.append(sub['Designed Layout'].sum())
                if 'Headcount' in sub.columns:
                    headcounts_val.append(sub['Headcount'].sum())
                else:
                    headcounts_val.append(0)

            if sum(jobs_val) > 0:
                report_lines.append("#### Output Trend (Past 5 Weeks):")
                report_lines.append("| Metrics | W-4 | W-3 | W-2 | W-1 | W0 (Current Week) |")
                report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: |")
                report_lines.append("| Completed Jobs | " + " | ".join([fmt_int(x) for x in jobs_val]) + " |")

                # Project Type sub-rows
                if len(comp_detail) > 0:
                    for jt in sorted(comp_detail['Job Type'].dropna().unique()):
                        jt_data = comp_detail[comp_detail['Job Type'] == jt]
                        jt_vals = []
                        for wl in week_labels:
                            sub_jt = jt_data[jt_data['Week'] == wl]
                            jt_vals.append(sub_jt['Jobs_Completed'].sum())
                        if sum(jt_vals) > 0:
                            report_lines.append(f"|     • {jt} | " + " | ".join([fmt_int(x) for x in jt_vals]) + " |")

                report_lines.append("| Headcount | " + " | ".join([fmt_int(x) for x in headcounts_val]) + " |")
                report_lines.append("| Dwelling Units | " + " | ".join([fmt_int(x) for x in units_val]) + " |")
                report_lines.append("| Area (Sqm) | " + " | ".join([fmt_sqm(x) for x in sqm_val]) + " |")
                report_lines.append("| Designed (Layouts) | " + " | ".join([fmt_int(x) for x in layouts_val]) + " |")
                report_lines.append("")

                # KPI inline
                actual_w0 = jobs_val[-1] if metric == "Jobs" else layouts_val[-1]
                if weekly_target > 0:
                    pct = (actual_w0 / weekly_target) * 100
                    status = "🟢 Passed" if pct >= 100 else "🔴 Missed"
                    report_lines.append(
                        f"**KPI W0:** Actual = **{fmt_int(actual_w0)}** {metric} / "
                        f"Target = {weekly_target} {metric} → **{pct:.1f}%** {status}"
                    )
                else:
                    report_lines.append(f"**KPI W0:** Actual = **{fmt_int(actual_w0)}** {metric} (No Target)")
            else:
                report_lines.append(f"*No completed jobs found for {company} in the past 5 weeks.*")

            # Leadtime for this company in W0
            if len(w0_completed_jobs) > 0:
                comp_lt = w0_completed_jobs[
                    (w0_completed_jobs['Team'] == team_name) &
                    (w0_completed_jobs['Company'] == company)
                ]
                if len(comp_lt) > 0:
                    avg_lt = comp_lt['Leadtime_Days'].mean()
                    report_lines.append(
                        f"**Leadtime W0:** Average **{fmt_float(avg_lt, decimals=2)} days** ({len(comp_lt)} jobs)"
                    )
            report_lines.append("")

            # ── Member performance scoped to this company ─────────────────────
            mc_members_df = mc_workload_df[
                (mc_workload_df['Company'] == company) &
                (mc_workload_df['userName'].isin(team_members))
            ]

            if len(mc_members_df) > 0:
                report_lines.append("#### Member Performance for this Client (W0):")
                report_lines.append(
                    "| Employee  | Task Design | Task Check | R&F | Training | "
                    "Dwelling Units) | Sqm | Designed (Layouts) | "
                    "Est.Design Median | Amd.Design Median |"
                )
                report_lines.append(
                    "| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |"
                )

                for _, mrow in mc_members_df.sort_values('Design', ascending=False).iterrows():
                    uname = mrow['userName']
                    # Complexity medians for this user × company
                    est_des = mc_complexity.get((uname, company, 'Est', 'Design'), None)
                    amd_des = mc_complexity.get((uname, company, 'Amd', 'Design'), None)

                    report_lines.append(
                        f"| {uname} "
                        f"| {fmt_int(mrow.get('Design', 0))} "
                        f"| {fmt_int(mrow.get('Check', 0))} "
                        f"| {fmt_int(mrow.get('R&F', 0))} "
                        f"| {fmt_int(mrow.get('Training', 0))} "
                        f"| {fmt_int(mrow.get('Design_Units', 0))} "
                        f"| {fmt_sqm(mrow.get('Design_Sqm', 0))} "
                        f"| {fmt_int(mrow.get('Design_Layouts', 0))} "
                        f"| {fmt_hrs(est_des)} "
                        f"| {fmt_hrs(amd_des)} |"
                    )
            else:
                report_lines.append("*No member completed tasks for this client in W0.*")

            report_lines.append("\n---\n")

        # ── TEAM SUMMARY SECTION ──────────────────────────────────────────────
        report_lines.append(f"## {t_idx}.{len(ordered_companies)+1}. Team Summary: {team_name}")
        report_lines.append("")

        # 5-week team total
        if len(team_trend) > 0 and team_trend['Jobs_Completed'].sum() > 0:
            report_lines.append("#### Total Team Workload (Past 5 Weeks - All Clients):")
            report_lines.append("| Metrics | W-4 | W-3 | W-2 | W-1 | W0 (Current week) |")
            report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: |")

            t_jobs = []; t_units = []; t_sqm = []; t_layouts = []
            for wl in week_labels:
                sub = team_trend[team_trend['Week'] == wl]
                t_jobs.append(sub['Jobs_Completed'].sum())
                t_units.append(sub['Dwelling Units'].sum())
                t_sqm.append(sub['Designed Square Meter'].sum())
                t_layouts.append(sub['Designed Layout'].sum())

            report_lines.append("| **Total Completed Jobs** | " + " | ".join([f"**{fmt_int(x)}**" for x in t_jobs]) + " |")
            report_lines.append("| Dwelling Units | " + " | ".join([fmt_int(x) for x in t_units]) + " |")
            report_lines.append("| Area (Sqm) | " + " | ".join([fmt_sqm(x) for x in t_sqm]) + " |")
            report_lines.append("| Designed (Layouts) | " + " | ".join([fmt_int(x) for x in t_layouts]) + " |")
            report_lines.append("")

        # Leadtime summary across all customers for W0
        if len(w0_completed_jobs) > 0:
            team_lt = w0_completed_jobs[w0_completed_jobs['Team'] == team_name]
            if len(team_lt) > 0:
                avg_lt_team = team_lt['Leadtime_Days'].mean()
                report_lines.append(f"**Team Average Leadtime (W0):** {fmt_float(avg_lt_team, decimals=2)} days ({len(team_lt)} jobs)")

                lt_by_comp = team_lt.groupby('Company')['Leadtime_Days'].agg(['count', 'mean']).reset_index()
                lt_by_comp.columns = ['Company', 'Jobs', 'Avg_LT']
                report_lines.append("\n**Average Leadtime by Client:**")
                report_lines.append("| Client | W0 Completed Jobs | Avg Leadtime (days) |")
                report_lines.append("| :--- | :---: | :---: |")
                for _, lr in lt_by_comp.sort_values('Avg_LT').iterrows():
                    report_lines.append(f"| {lr['Company']} | {fmt_int(lr['Jobs'])} | {fmt_float(lr['Avg_LT'], decimals=2)} |")
                report_lines.append("")

        # Member productivity summary (all customers, W0)
        team_workload = member_workload[member_workload['userName'].isin(team_members)]
        if len(team_workload) > 0:
            report_lines.append("#### Total Member Output W0 (All Clients):")
            report_lines.append(
                "| Employee | Design | Check | R&F | Training | "
                "Dwelling Units | Sqm | Designed Layouts |"
            )
            report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: |")
            for _, row in team_workload.sort_values('Design', ascending=False).iterrows():
                report_lines.append(
                    f"| {row['userName']} "
                    f"| {fmt_int(row['Design'])} "
                    f"| {fmt_int(row['Check'])} "
                    f"| {fmt_int(row['R&F'])} "
                    f"| {fmt_int(row['Training'])} "
                    f"| {fmt_int(row['Design_Units'])} "
                    f"| {fmt_sqm(row['Design_Sqm'])} "
                    f"| {fmt_int(row['Design_Layouts'])} |"
                )
            report_lines.append("")

        # Intraweek tasks
        team_intraweek = intraweek_summary[intraweek_summary['userName'].isin(team_members)]
        if len(team_intraweek) > 0:
            report_lines.append("#### Tasks Started & Completed in Same Week (Intraweek):")
            report_lines.append("| Member | Design | Check | R&F | Training | Total Intraweek |")
            report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: |")
            for _, row in team_intraweek.iterrows():
                tot = row['Design'] + row['Check'] + row['R&F'] + row['Training']
                report_lines.append(
                    f"| {row['userName']} "
                    f"| {fmt_int(row['Design'])} "
                    f"| {fmt_int(row['Check'])} "
                    f"| {fmt_int(row['R&F'])} "
                    f"| {fmt_int(row['Training'])} "
                    f"| {fmt_int(tot)} |"
                )
            report_lines.append("")

        # Attendance Quality (all team members)
        if len(attendance_summary) > 0:
            team_att = attendance_summary[attendance_summary['Workload_User'].isin(team_members)]
            if len(team_att) > 0:
                report_lines.append("#### Attendance & Working Hours:")
                report_lines.append(
                    "| Employee | Employee ID | Working Days | Standard Working Hour | Actual Working Time | OT | Difference | "
                    "Late CI (>10m) | Early CO | Leave | Unpaid Leave |"
                )
                report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: | :---: |")
                for _, row in team_att.sort_values('Actual_Hours', ascending=False).iterrows():
                    diff = row['Actual_Hours'] - row['Standard_Hours']
                    report_lines.append(
                        f"| {row['Workload_User']} | {row['EmployeeID']} "
                        f"| {fmt_float(row.get('Working_Days', 0.0), decimals=2)} | {fmt_float(row['Standard_Hours'], suffix='h', decimals=1)} | {fmt_float(row['Actual_Hours'], suffix='h', decimals=1)} | {fmt_float(row.get('OT_Hours', 0), suffix='h', decimals=1)} | {fmt_float(diff, suffix='h', decimals=2, show_plus=True)} "
                        f"| {fmt_int(row['Late_Checkins'])} | {fmt_int(row['Early_Checkouts'])} "
                        f"| {row['Annual_Leave'] if row['Annual_Leave'] > 0 else '-'} "
                        f"| {row['Unpaid_Leave'] if row['Unpaid_Leave'] > 0 else '-'} |"
                    )
            else:
                report_lines.append("\n*No attendance data found for team members.*")
        else:
            report_lines.append("\n*Unable to load or process attendance data.*")

        report_lines.append("\n---\n")

    # ── METRIC 7: JOB BACKLOG & AGING PROFILE ────────────────────────────────────
    print("Calculating Job Backlog & Aging Profile...")
    w0_end_naive = w0_end.replace(tzinfo=None)
    
    backlog_df = jobs_df[
        (jobs_df['Status'] != 'Complete') &
        (jobs_df['Time Created'].notna())
    ].copy()
    
    backlog_df['Time Created_Naive'] = backlog_df['Time Created'].dt.tz_localize(None)
    backlog_df = backlog_df[backlog_df['Time Created_Naive'] <= w0_end_naive]
    
    backlog_df['Aging_Days'] = (w0_end_naive - backlog_df['Time Created_Naive']).dt.days
    backlog_df['Team'] = backlog_df['Job #'].apply(attribute_job_to_team)
    
    def get_aging_bucket(days):
        if days <= 1:
            return '0-1 days'
        elif days == 2:
            return '2 days'
        elif days == 3:
            return '3 days'
        elif days == 4:
            return '4 days'
        elif days == 5:
            return '5 days'
        else:
            return '>5 days'
            
    backlog_df['Aging_Bucket'] = backlog_df['Aging_Days'].apply(get_aging_bucket)
    
    # Compile Backlog Table
    report_lines.append("# 6. JOB BACKLOG & AGING PROFILE REPORT")
    report_lines.append(f"Summary table of uncompleted workload (Job Backlog) as of reporting week end ({w0_end.strftime('%d/%m/%Y')}), categorized by aging profile:")
    report_lines.append("")
    report_lines.append("| Team / Client | 0-1 Days | 2 Days | 3 Days | 4 Days | 5 Days | >5 Days | Total Backlog |")
    report_lines.append("| :--- | :---: | :---: | :---: | :---: | :---: | :---: | :---: |")
    
    buckets = ['0-1 days', '2 days', '3 days', '4 days', '5 days', '>5 days']
    team_list = ['Frame & Truss', 'Drafting', 'Estimating', 'Engineered Wood Products (EWP)', 'Cleveland', 'Unknown']
    
    overall_totals = {b: 0 for b in buckets}
    overall_totals['Total'] = 0
    
    for t_name in team_list:
        t_df = backlog_df[backlog_df['Team'] == t_name]
        if len(t_df) == 0:
            continue
            
        t_label = t_name.upper() if t_name != 'Unknown' else 'CHƯA PHÂN LOẠI'
        
        # Calculate Team Totals
        t_totals = {}
        t_tot_sum = 0
        for b in buckets:
            c = len(t_df[t_df['Aging_Bucket'] == b])
            t_totals[b] = c
            t_tot_sum += c
            overall_totals[b] += c
            
        overall_totals['Total'] += t_tot_sum
        
        report_lines.append(f"| **{t_label}** | " + " | ".join([f"**{fmt_int(t_totals[b])}**" for b in buckets]) + f" | **{fmt_int(t_tot_sum)}** |")
        
        # Companies under this team
        companies = sorted(list(t_df['Company'].dropna().unique()))
        for comp in companies:
            c_df = t_df[t_df['Company'] == comp]
            c_totals = {}
            c_tot_sum = 0
            for b in buckets:
                c = len(c_df[c_df['Aging_Bucket'] == b])
                c_totals[b] = c
                c_tot_sum += c
                
            report_lines.append(f"| - {comp} | " + " | ".join([fmt_int(c_totals[b]) for b in buckets]) + f" | {fmt_int(c_tot_sum)} |")
            
    # Add Overall Total Row
    report_lines.append(f"| **TOTAL** | " + " | ".join([f"**{fmt_int(overall_totals[b])}**" for b in buckets]) + f" | **{fmt_int(overall_totals['Total'])}** |")
    report_lines.append("")
    report_lines.append("---\n")

    # ── METRIC 8 (Section 7): TASK COMPLEXITY PROFILE ────────────────────────────
    print("Calculating Task Complexity Profile (investedSeconds)...")

    # Parse timeOutcome using the robust helper
    tasks_df['timeOutcome_parsed'] = parse_date_column(tasks_df['timeOutcome'])
    tasks_df['timeOutcome_naive'] = tasks_df['timeOutcome_parsed'].apply(
        lambda x: x.replace(tzinfo=None) if pd.notna(x) else pd.NaT
    )

    w0_start_naive = w0_start.replace(tzinfo=None)
    w0_end_naive2  = w0_end.replace(tzinfo=None)

    # Filter: completed in week, has investedSeconds > 0
    week_tasks = tasks_df[
        (tasks_df['Flow Task_status'] == 'Complete') &
        (tasks_df['timeOutcome_naive'].notna()) &
        (tasks_df['timeOutcome_naive'] >= w0_start_naive) &
        (tasks_df['timeOutcome_naive'] <= w0_end_naive2) &
        (tasks_df['investedSeconds'].notna()) &
        (tasks_df['investedSeconds'] > 0)
    ].copy()

    # Derive Team from userName
    week_tasks['Team'] = week_tasks['userName'].map(user_to_team).fillna('Unknown')
    # Exclude system/admin users and unknowns that are excluded
    week_tasks = week_tasks[~week_tasks['Team'].isin(['Unknown', 'Excluded'])]

    # Extract Project Prefix (Est / Amd / Det / Admin / Other)
    def get_project_prefix(s):
        s = str(s).strip()
        if s.startswith('Est.'):   return 'Est'
        if s.startswith('Amd.'):   return 'Amd'
        if s.startswith('Det.'):   return 'Det'
        if s.lower().startswith('admin'): return 'Admin'
        if s.lower().startswith('training') or s.lower().startswith('adhoc'): return 'Training/Adhoc'
        return 'Other'

    # Extract Flow Type
    def get_flow_type(s):
        sl = str(s).lower()
        if 'design'  in sl: return 'Design'
        if 'check'   in sl: return 'Check'
        if 'review'  in sl and 'fix' in sl: return 'R&F'
        if 'training' in sl: return 'Training'
        if 'admin'   in sl or 'qs' in sl: return 'Admin-QS'
        return 'Other'

    week_tasks['ProjectPrefix'] = week_tasks['taskTypeName'].apply(get_project_prefix)
    week_tasks['FlowType']      = week_tasks['taskTypeName'].apply(get_flow_type)
    week_tasks['investedHours'] = week_tasks['investedSeconds'] / 3600.0

    # Utility: format hours as "Xh Ym" or "-"
    def fmt_hrs(val):
        if pd.isna(val) or val == 0:
            return '-'
        h = int(val)
        m = int(round((val - h) * 60))
        if h > 0 and m > 0:  return f"{h}h{m:02d}m"
        if h > 0:             return f"{h}h"
        return f"{m}m"

    # ─── 7.1 Team-level Summary ────────────────────────────────────────────────
    report_lines.append("# 7. TASK COMPLEXITY PROFILE — AVERAGE HOUR PER TASK")
    report_lines.append(
        f"Analyzes the actual effort spent (`investedSeconds`) on tasks completed "
        f"during the week ({w0_start.strftime('%d/%m/%Y')} – {w0_end.strftime('%d/%m/%Y')}). "
        "GHelps assess the workload associated with each project type (Est/Amd/Det) and "
        "each work phase (Design / Check / R&F)."
    )
    report_lines.append("")

    # Pivot prefixes we care about (exclude Admin/Other for complexity read)
    COMPLEXITY_PREFIXES = ['Est', 'Amd', 'Det']
    FLOW_COLS = ['Design', 'Check', 'R&F']

    # Build header
    header_parts = ["| Team | Project | # Tasks | Median (hrs) | Mean (hrs) |"]
    sep_parts    = ["| :--- | :--- | ---: | ---: | ---: |"]
    report_lines += header_parts + sep_parts

    active_teams_ordered = ['Frame & Truss', 'Drafting', 'Estimating',
                             'Engineered Wood Products (EWP)', 'Cleveland']

    for t_name in active_teams_ordered:
        t_df = week_tasks[week_tasks['Team'] == t_name]
        if len(t_df) == 0:
            continue

        for prefix in COMPLEXITY_PREFIXES:
            p_df = t_df[t_df['ProjectPrefix'] == prefix]
            if len(p_df) == 0:
                continue

            # Team-prefix row header
            report_lines.append(f"| **{t_name}** | **{prefix}.** | | | |")

            for flow in FLOW_COLS:
                f_df = p_df[p_df['FlowType'] == flow]
                if len(f_df) == 0:
                    continue
                count  = len(f_df)
                median = f_df['investedHours'].median()
                mean   = f_df['investedHours'].mean()
                report_lines.append(
                    f"|    ↳ | {flow} | {count} | {fmt_hrs(median)} | {fmt_hrs(mean)} |"
                )

    report_lines.append("")

    # ─── 7.2 Member Drill-Down ────────────────────────────────────────────────
    report_lines.append("## 7.2 Member Breakdown")
    report_lines.append(
        "Shows the median time spent by each team member on Design and Check tasks "
        "across different project types during the week:"
    )
    report_lines.append("")

    # Header: Member | Est.Design | Est.Check | Amd.Design | Amd.Check | Det.Design | Det.Check | Total Tasks
    col_specs = [
        ('Est', 'Design'),  ('Est', 'Check'),
        ('Amd', 'Design'),  ('Amd', 'Check'),
        ('Det', 'Design'),  ('Det', 'Check'),
    ]
    col_headers = " | ".join([f"{p}.{f[:3]}" for (p, f) in col_specs])
    report_lines.append(
        f"| Team | Member | {col_headers} | # Tasks |"
    )
    report_lines.append(
        "| :--- | :--- | " + " | ".join(["---:"] * (len(col_specs) + 1)) + " |"
    )

    for t_name in active_teams_ordered:
        t_df = week_tasks[week_tasks['Team'] == t_name]
        if len(t_df) == 0:
            continue

        members = sorted(t_df['userName'].unique())
        first_member = True
        for member in members:
            m_df = t_df[t_df['userName'] == member]
            total_tasks = len(m_df)
            values = []
            for (prefix, flow) in col_specs:
                cell_df = m_df[(m_df['ProjectPrefix'] == prefix) & (m_df['FlowType'] == flow)]
                if len(cell_df) == 0:
                    values.append('-')
                else:
                    values.append(fmt_hrs(cell_df['investedHours'].median()))

            team_label = f"**{t_name}**" if first_member else ""
            first_member = False
            report_lines.append(
                f"| {team_label} | {member} | " + " | ".join(values) + f" | {total_tasks} |"
            )

    report_lines.append("")
    report_lines.append("---\n")
    report_lines.append("*(Report automatically generated by Weekly HR Report Framework)*")

    # Save report
    os.makedirs(output_dir, exist_ok=True)
    report_file_name = f"Weekly_Performance_Report_{w0_end.strftime('%Y%m%d')}.docx"
    report_file_path = os.path.join(output_dir, report_file_name)
    
    try:
        final_path = generate_docx_report(report_lines, report_file_path)
        print(f"\nWeekly HR report generated successfully and saved at: {final_path}")
        
        # Export dashboard JSON
        import subprocess
        temp_md_path = report_file_path.replace('.docx', '.tmp.md')
        with open(temp_md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_lines))

        json_out_path = os.path.join(output_dir, "dashboard", "data.js")
        os.makedirs(os.path.dirname(json_out_path), exist_ok=True)
        parser_script = os.path.join(os.path.dirname(__file__), "parse_dashboard_json.py")
        if os.path.exists(parser_script):
            subprocess.run([sys.executable, parser_script, temp_md_path, json_out_path])

        if os.path.exists(temp_md_path):
            os.remove(temp_md_path)
            
    except Exception as e:
        print(f"Error saving report file: {e}")

if __name__ == "__main__":
    main()
